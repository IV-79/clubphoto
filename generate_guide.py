# -*- coding: utf-8 -*-
"""
Genere guide_du_site.docx - Club Photo Tain-Tournon
Usage : python generate_guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def h1(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def para(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def note(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFBEB')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    run = p.add_run(u'ℹ  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x78, 0x35, 0x0f)
    doc.add_paragraph()
    return tbl

def fields_table(rows):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, desc) in enumerate(rows):
        row = tbl.rows[i]
        c0 = row.cells[0]
        c0.width = Cm(5)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.name = 'Calibri'
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
    doc.add_paragraph()
    return tbl

STATUT_COLORS = {
    'preparation': ('F3F4F6', '374151'),
    'ideation':    ('E0F2FE', '075985'),
    'nettoyage':   ('FCE7F3', '9D174D'),
    'vote':        ('FEF3C7', '92400E'),
    'votation':    ('FEF3C7', '92400E'),
    'ouvert':      ('DCFCE7', '166534'),
    'soumission':  ('DCFCE7', '166534'),
    'resultats':   ('EDE9FE', '5B21B6'),
    'cloture':     ('EDE9FE', '5B21B6'),
    'inscription': ('DBEAFE', '1D4ED8'),
    'fermeture':   ('E0F2FE', '0369A1'),
    'a_venir':     ('F3F4F6', '374151'),
    'en_cours':    ('DCFCE7', '166534'),
}

def statuts_table(rows):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (statut, desc, key) in enumerate(rows):
        bg, fg = STATUT_COLORS.get(key, ('F3F4F6', '374151'))
        row = tbl.rows[i]
        c0 = row.cells[0]
        c0.width = Cm(4.5)
        tc0 = c0._tc
        tcPr0 = tc0.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr0.append(shd)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(statut)
        r0.bold = True
        r0.font.name = 'Calibri'
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor(int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
    doc.add_paragraph()
    return tbl

def separator():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D1D5DB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def page_break():
    doc.add_page_break()


# ===========================================================================
# PAGE DE TITRE
# ===========================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Club Photo Tain-Tournon')
run.font.name = 'Calibri'; run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guide du site')
run.font.name = 'Calibri'; run.font.size = Pt(28); run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    u"A destination de tous les membres — consultez ce guide pour "
    u"d\xe9couvrir le fonctionnement du site et de toutes ses fonctionnalit\xe9s."
)
run.font.name = 'Calibri'; run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
separator()
doc.add_paragraph()


# ===========================================================================
# 1. INFORMATIONS GENERALES
# ===========================================================================
h1(u'1. Informations g\xe9n\xe9rales')

h2(u'1.1 A quoi sert ce site ?')
para(
    u"Le site du Club Photo Tain-Tournon centralise tout ce dont les membres ont besoin au "
    u"quotidien : suivre les actualit\xe9s, consulter le calendrier des activit\xe9s, "
    u"partager et d\xe9couvrir des photos, participer aux comp\xe9titions et retrouver les "
    u"documents du club."
)
para(
    u"Une partie du site est accessible sans connexion — actualit\xe9s publiques, "
    u"portfolios des membres, galeries. Les fonctions interactives (voter, commenter, "
    u"s’inscrire \xe0 une activit\xe9, uploader des photos) n\xe9cessitent "
    u"d’\xeatre connect\xe9 avec son compte membre."
)

h3(u'Ce qui est g\xe9r\xe9 sur le site')
fields_table([
    (u'Actualit\xe9s',
     u'Les nouvelles du club : infos, annonces, comptes-rendus publi\xe9s par les contributeurs et admins.'),
    (u'Calendrier',
     u'Toutes les dates importantes. Accessible sans connexion ; les r\xe9unions du club et les '
     u'activit\xe9s « Membres seulement » sont masqu\xe9es pour les visiteurs non connect\xe9s.'),
    (u'Portfolios',
     u'Chaque membre dispose d’un espace photo personnel. Les visiteurs peuvent parcourir les portfolios publics.'),
    (u'Documents',
     u'Statuts, comptes-rendus, formulaires — un espace de partage r\xe9serv\xe9 aux membres connect\xe9s.'),
    (u'Th\xe8me du mois',
     u'La comp\xe9tition mensuelle du club : soumission de photos sur un th\xe8me impos\xe9, puis vote.'),
    (u'Activit\xe9s',
     u'Sorties, ateliers, d\xe9fis photo, OneShots et expositions — organis\xe9s par les contributeurs et admins.'),
])

h2(u'1.2 Les types d’activit\xe9s')
para(u'Le site g\xe8re six types d’activit\xe9s, chacun avec son propre fonctionnement :')
fields_table([
    (u'Sortie photo',
     u'On se retrouve \xe0 un endroit pour photographier ensemble. Inscription optionnelle, '
     u'chacun uploade ses propres photos apr\xe8s la sortie.'),
    (u'Sortie club',
     u'Balade, visite de mus\xe9e, sortie culturelle — pas forc\xe9ment centr\xe9e sur la photo. '
     u'M\xeame fonctionnement que la sortie photo.'),
    (u'D\xe9fi photo',
     u'Comp\xe9tition th\xe9matique. Les membres s’inscrivent, uploadent leurs photos, puis '
     u'votent anonymement. Les phases changent automatiquement selon les dates.'),
    (u'Atelier',
     u"Session d’apprentissage ou de travail en groupe (retouche, technique, partage "
     u"d’exp\xe9rience). Pas de vote, pas de comp\xe9tition."),
    (u'OneShot',
     u"\xc9v\xe9nement photo comp\xe9titif avec th\xe8mes multiples g\xe9r\xe9 manuellement par "
     u"l’organisateur. C’est l’organisateur qui uploade les photos des participants."),
    (u'Exposition',
     u"Projet d’exposition photo en plusieurs phases : id\xe9ation collective des th\xe8mes, "
     u"vote, soumission des photos, puis cl\xf4ture. Nouveau type d’activit\xe9 sur le site."),
])

h2(u'1.3 Comment naviguer dans le site')
para(
    u"La navigation suit toujours le m\xeame sch\xe9ma : page liste (grille de cartes) "
    u"→ fiche info (d\xe9tails complets) → plein \xe9cran (photos uniquement)."
)
bullet(u"Depuis une page liste, cliquez sur une carte pour ouvrir la fiche info.")
bullet(u"Dans la fiche info, cliquez sur une photo pour l’ouvrir en plein \xe9cran.")
bullet(u"Naviguez entre les photos avec les fl\xe8ches \xe0 l’\xe9cran, les touches ← → du clavier, ou en glissant horizontalement sur mobile.")
note(
    u"Pendant les votes (Th\xe8me du mois, OneShot, D\xe9fi), un clic ouvre directement le plein \xe9cran "
    u"sans nom d’auteur ni commentaires — les soumissions sont anonymes."
)

h2(u'1.4 Uploader une photo')
para(u"Le site accepte les fichiers JPEG et PNG. Chaque photo est automatiquement compress\xe9e c\xf4t\xe9 client avant l’envoi.")
note(
    u"Uploadez toujours le fichier original sorti de l’appareil ou du logiciel de retouche — "
    u"pas une capture d’\xe9cran, pas une photo re\xe7ue par messagerie. WhatsApp et Facebook "
    u"suppriment les m\xe9tadonn\xe9es EXIF."
)

h2(u'1.5 M\xe9tadonn\xe9es EXIF')
para(u"Quand vous uploadez une photo, le site lit automatiquement les m\xe9tadonn\xe9es EXIF et les affiche dans la fiche info :")
fields_table([
    (u'Appareil',            u'Marque et mod\xe8le (ex : Sony A7 IV).'),
    (u'Objectif',            u'Objectif utilis\xe9 (ex : 85mm f/1.8).'),
    (u'Focale',              u'Focale effective en millim\xe8tres.'),
    (u'Ouverture',           u'Valeur f/ utilis\xe9e (ex : f/2.8).'),
    (u'Vitesse',             u"Temps d’exposition (ex : 1/500 s)."),
    (u'ISO',                 u'Sensibilit\xe9 ISO au moment de la prise de vue.'),
    (u'Date de prise de vue',u"Date enregistr\xe9e par l’appareil, pas la date d’upload."),
    (u'D\xe9finition',       u"R\xe9solution en pixels et m\xe9gapixels (ex : 6 000 \xd7 4 000 px — 24 Mpx). Calcul\xe9e \xe0 partir de l’image elle-m\xeame."),
    (u'Coordonn\xe9es GPS',  u'Lieu de prise de vue. Affich\xe9 comme lien « Voir sur la carte » dans la fiche info.'),
])
note(
    u"Vie priv\xe9e GPS : si votre photo contient des coordonn\xe9es GPS, une fen\xeatre vous demande "
    u"si vous souhaitez les conserver. Par d\xe9faut Oui — utile pour les paysages. "
    u"R\xe9pondez Non pour une photo prise chez vous."
)

h2(u'1.6 Likes ♥ et commentaires')
para(
    u"Chaque photo peut \xeatre lik\xe9e depuis la carte ou la fiche info. Un cœur rouge "
    u"signifie que vous avez lik\xe9 cette photo — cliquez \xe0 nouveau pour retirer votre like. "
    u"Les commentaires sont accessibles depuis la fiche info uniquement. Connexion requise."
)

page_break()


# ===========================================================================
# 2. POUR LES MEMBRES
# ===========================================================================
h1(u'2. Pour les membres')
para(u"Cette section d\xe9crit les actions disponibles pour tout membre connect\xe9.")

h2(u'2.1 Actualit\xe9s')
para(
    u"La page Actualit\xe9s affiche toutes les publications sous forme de cartes. Les articles "
    u"\xe9pingl\xe9s apparaissent en t\xeate. Filtrez par type avec les boutons en haut de page. "
    u"Cliquez sur une carte pour lire l’article complet."
)

h2(u'2.2 Calendrier')
para(
    u"Le Calendrier regroupe les r\xe9unions, les \xe9v\xe9nements et les activit\xe9s \xe0 venir. "
    u"Cliquez sur un \xe9l\xe9ment pour voir les d\xe9tails. Des filtres permettent d’afficher "
    u"un seul type \xe0 la fois."
)
note(
    u"Accessible sans connexion. Les r\xe9unions du club et les activit\xe9s \xe0 visibilit\xe9 "
    u"« Membres seulement » n’apparaissent pas pour les visiteurs non connect\xe9s."
)

h2(u'2.3 Documents')
h3(u'Consulter et t\xe9l\xe9charger')
para(u"Accessible via Membres → Documents. Les fichiers sont class\xe9s par dossiers. Filtrez par dossier avec les boutons en haut, puis cliquez ⬇ pour ouvrir ou t\xe9l\xe9charger.")
h3(u'D\xe9poser un document')
numbered(u"Cliquez Nouveau document.")
numbered(u"Renseignez le nom et choisissez le dossier de destination.")
numbered(u"Glissez votre fichier ou cliquez pour le s\xe9lectionner, puis Uploader.")
para(u"Vous pouvez remplacer (ic\xf4ne crayon) ou supprimer (ic\xf4ne poubelle) les documents que vous avez d\xe9pos\xe9s.")

h2(u'2.4 Mon portfolio')
para(
    u"Depuis Mon espace → Mon Portfolio, uploadez vos photos, organisez-les par cat\xe9gories "
    u"et choisissez leur visibilit\xe9 (Public ou Membres uniquement). Votre portfolio appara\xeet "
    u"dans la galerie des membres d\xe8s que vous avez au moins une photo publique."
)

h2(u'2.5 Th\xe8me du mois')
h3(u'Le principe')
para(
    u"Chaque mois, le club propose un th\xe8me photo. Les membres soumettent leurs photos, "
    u"puis votent pour leurs pr\xe9f\xe9r\xe9es. Les phases changent automatiquement selon les dates."
)
statuts_table([
    (u'En attente', u'Avant le 1er du mois — le th\xe8me est annonc\xe9.', 'a_venir'),
    (u'Ouvert',     u'Tout le mois — soumettez vos photos.',               'ouvert'),
    (u'Vote',       u'Apr\xe8s la cl\xf4ture — votez pour vos pr\xe9f\xe9r\xe9es.', 'vote'),
    (u'R\xe9sultats',u'Apr\xe8s la fin du vote — podium et r\xe9v\xe9lation.', 'resultats'),
])
h3(u'Soumettre une photo')
numbered(u"Ouvrez le th\xe8me dont le statut est Ouvert.")
numbered(u"Glissez votre photo dans la zone d’upload ou cliquez pour choisir un fichier.")
numbered(u"Pour retirer votre soumission, cliquez \U0001f5d1 sur votre photo avant la cl\xf4ture.")
note(u"Vous ne pouvez soumettre qu’une photo par th\xe8me (sauf indication contraire de l’organisateur).")
h3(u'Voter')
numbered(u"Pendant la phase Vote, les photos sont anonymes.")
numbered(u"Cliquez Voter sous les photos que vous pr\xe9f\xe9rez. Vous avez un nombre limit\xe9 de votes.")
numbered(u"Cliquez ✓ Retirer pour reprendre un vote et le d\xe9poser ailleurs.")
note(u"Vous ne pouvez pas voter pour votre propre photo.")

h2(u'2.6 Sorties photo, Sorties club & Ateliers')
numbered(u"Ouvrez l’\xe9v\xe9nement depuis Activit\xe9s → \xc9v\xe9nements.")
numbered(u"Cliquez Je participe pour vous inscrire.")
numbered(u"Apr\xe8s l’\xe9v\xe9nement, revenez sur la page et uploadez vos photos. Plusieurs fichiers accept\xe9s en une fois.")
note(u"La photo la plus lik\xe9e devient automatiquement la photo de couverture de l’\xe9v\xe9nement.")

h2(u'2.7 D\xe9fi Photo')
h3(u'Le principe')
para(
    u"Un organisateur lance un d\xe9fi sur un th\xe8me pr\xe9cis. Les membres s’inscrivent, "
    u"uploadent leurs propres photos pendant la p\xe9riode de soumission, puis votent. "
    u"Les phases changent automatiquement selon les dates."
)
statuts_table([
    (u'A venir',    u'Annonc\xe9, inscriptions ouvertes.',                                    'a_venir'),
    (u'En cours',   u'Uploadez vos photos. Vous ne voyez que les v\xf4tres (anonymat).',      'en_cours'),
    (u'Vote',       u'Toutes les photos r\xe9v\xe9l\xe9es, noms masqu\xe9s. Votez.',         'vote'),
    (u'R\xe9sultats',u'Classement et podium avec r\xe9v\xe9lation des auteurs.',              'resultats'),
])
h3(u'Participer')
numbered(u"Ouvrez un d\xe9fi et cliquez S’inscrire.")
numbered(u"Pendant la phase En cours, uploadez vos photos dans la zone pr\xe9vue.")
numbered(u"Pendant le Vote, parcourez toutes les photos et votez.")
note(u"Seuls les membres inscrits peuvent soumettre des photos et voter.")

h2(u'2.8 OneShot')
h3(u'Le principe')
para(
    u"Un OneShot est organis\xe9 sur un ou plusieurs th\xe8mes. Vous vous inscrivez, "
    u"l’organisateur uploade vos photos \xe0 votre place, puis ouvre le vote quand tout est pr\xeat."
)
statuts_table([
    (u'Inscriptions ouvertes', u'Inscrivez-vous librement.',                                                   'inscription'),
    (u'Inscriptions ferm\xe9es',u"Plus d’inscription possible, l’organisateur finalise les photos.", 'fermeture'),
    (u'Vote',                  u'Th\xe8mes et photos r\xe9v\xe9l\xe9s. Votez pour une photo par th\xe8me, pas la v\xf4tre.', 'vote'),
    (u'R\xe9sultats',          u'Classement publi\xe9 avec r\xe9v\xe9lation des auteurs.',                    'resultats'),
])
h3(u'Participer')
numbered(u"Cliquez Je participe pendant la phase Inscriptions ouvertes.")
numbered(u"L’organisateur s’occupe d’uploader vos photos.")
numbered(u"Pendant le Vote, votez pour une photo par th\xe8me (pas la v\xf4tre). Les noms sont masqu\xe9s.")

h2(u'2.9 Exposition')
h3(u'Le principe')
para(u"Une exposition se d\xe9roule en plusieurs phases. Les membres participent \xe0 la d\xe9finition du th\xe8me, puis soumettent leurs photos.")
statuts_table([
    (u'Id\xe9ation',  u"Proposez des id\xe9es de th\xe8mes pour l’exposition.",          'ideation'),
    (u'Vote',         u'Votez pour les th\xe8mes qui vous inspirent le plus.',                'vote'),
    (u'Soumission',   u"Le th\xe8me est choisi. Uploadez vos photos pour l’exposition.",'soumission'),
    (u'Cl\xf4tur\xe9',u"La galerie des photos de l’exposition est accessible.",          'cloture'),
])
h3(u'Participer')
bullet(u"Id\xe9ation : proposez un ou plusieurs th\xe8mes dans la zone pr\xe9vue.")
bullet(u"Vote : utilisez vos jetons de vote pour soutenir les th\xe8mes. Vous pouvez r\xe9partir vos votes sur plusieurs th\xe8mes.")
bullet(u"Soumission : uploadez vos photos dans la zone d\xe9di\xe9e, dans la limite fix\xe9e par l’organisateur.")

page_break()


# ===========================================================================
# 3. POUR LES CONTRIBUTEURS
# ===========================================================================
h1(u'3. Pour les contributeurs')
para(
    u"Un contributeur peut organiser toutes les activit\xe9s du club, publier des articles, "
    u"g\xe9rer les th\xe8mes du mois et les r\xe9unions. Pour devenir contributeur, "
    u"contactez un administrateur."
)

h2(u'3.1 Th\xe8mes du mois')
para(u"Accessible via le bouton G\xe9rer les th\xe8mes du mois en haut de la page Th\xe8mes du mois.")
fields_table([
    (u'Titre + mois',    u'Le mois d\xe9termine automatiquement les 3 dates : ouverture (1er du mois), cl\xf4ture (dernier jour), fin du vote (cl\xf4ture + N jours).'),
    (u'Jours de vote',   u'D\xe9faut : 15. Modifiable jusqu’au d\xe9but du vote.'),
    (u'Photos / membre', u'D\xe9faut : 1.'),
    (u'Votes / membre',  u'D\xe9faut : 3.'),
])
note(
    u"Il n’y a pas de bouton « passer \xe0 l’\xe9tape suivante » — "
    u"le statut est calcul\xe9 automatiquement. Pour changer le mois, supprimez et recr\xe9ez le th\xe8me "
    u"(possible tant qu’aucune soumission n’existe)."
)

h2(u'3.2 R\xe9unions')
para(u"Accessible via le bouton G\xe9rer les r\xe9unions en haut du Calendrier.")
para(
    u"Cr\xe9ez, modifiez et supprimez les r\xe9unions du club (assembl\xe9es, conseils "
    u"d’administration, s\xe9ances de travail). Une notification est envoy\xe9e aux membres "
    u"\xe0 la cr\xe9ation et si la date change."
)

h2(u'3.3 Publier un article')
para(u"Le bouton + Nouvel article est visible en haut de la page Actualit\xe9s.")
fields_table([
    (u'Type',     u'Info, Encart, Exposition, \xc9v\xe9nement, Annonce — d\xe9termine la couleur du badge.'),
    (u'Statut',   u'Brouillon (non visible) ou Publi\xe9. Un article publi\xe9 est imm\xe9diatement visible selon la port\xe9e choisie.'),
    (u'Port\xe9e',u'Public (tout le monde) ou Membres uniquement.'),
    (u'\xc9pingler',u"L’article reste en t\xeate de liste quelle que soit sa date."),
])
note(u"Pour modifier un article depuis la page Actualit\xe9s, cliquez sur l’ic\xf4ne crayon ✎ en haut \xe0 droite de sa carte.")

h2(u'3.4 Cr\xe9er une activit\xe9')
para(u"Allez sur Activit\xe9s → \xc9v\xe9nements et cliquez + Organiser un \xe9v\xe9nement. Choisissez ensuite le type d’activit\xe9.")

h2(u'3.5 Sorties & Ateliers')
h3(u'Cr\xe9er')
numbered(u"Choisissez le type : Sortie Photo, Sortie Club ou Atelier.")
numbered(u"Renseignez titre, date, lieu et description.")
numbered(u"Choisissez la visibilit\xe9 : Public ou Membres seulement.")
numbered(u"Cochez Inscription obligatoire si vous souhaitez compter les participants.")
numbered(u"L’\xe9v\xe9nement est visible imm\xe9diatement apr\xe8s cr\xe9ation.")
h3(u'G\xe9rer')
bullet(u"Modifiez le titre, la date, le lieu et la description \xe0 tout moment via ✏ Modifier.")
bullet(u"G\xe9rez la liste des participants dans la section Inscriptions.")
bullet(u"Vous pouvez supprimer les photos upload\xe9es par les participants.")
bullet(u"Pour supprimer l’\xe9v\xe9nement : bouton \U0001f5d1 Supprimer en haut de la page.")

h2(u'3.6 D\xe9fi Photo')
h3(u'Cr\xe9er')
numbered(u"S\xe9lectionnez le type D\xe9fi Photo.")
numbered(u"Renseignez le titre et le th\xe8me \xe0 photographier (obligatoire).")
numbered(u"D\xe9finissez les 3 dates : d\xe9but des soumissions, fin des soumissions, cl\xf4ture des votes.")
numbered(u"Choisissez le max de photos par participant et le nombre de votes par membre.")
numbered(u"Cliquez Cr\xe9er le d\xe9fi — une notification est envoy\xe9e aux membres.")
note(u"Les phases changent automatiquement selon les dates — aucune intervention manuelle n\xe9cessaire.")
h3(u'G\xe9rer')
bullet(u"Modifiez les dates, le titre et la description via ✏ Modifier tant que le d\xe9fi n’est pas en phase Vote.")
bullet(u"G\xe9rez les inscriptions et supprimez des photos si n\xe9cessaire.")

h2(u'3.7 OneShot')
h3(u'Cr\xe9er')
numbered(u"S\xe9lectionnez le type OneShot.")
numbered(u"Renseignez le titre (obligatoire). Date et lieu sont optionnels.")
numbered(u"Cliquez Cr\xe9er — vous arrivez sur la page en mode Pr\xe9paration (invisible aux autres membres).")
numbered(u"Dans la section Gestion, ajoutez vos th\xe8mes.")
numbered(u"Quand tout est pr\xeat, cliquez Ouvrir les inscriptions.")
h3(u'Cycle de vie')
statuts_table([
    (u'Pr\xe9paration',          u'Configurez titre, description, th\xe8mes. Invisible aux membres.',                    'preparation'),
    (u'Inscriptions ouvertes',   u"Visible. Les membres s’inscrivent. G\xe9rez les photos via G\xe9rer les photos.", 'inscription'),
    (u'Inscriptions ferm\xe9es', u"Plus d’inscription autonome. Finalisez les photos.",                             'fermeture'),
    (u'Vote',                    u"Th\xe8mes et photos r\xe9v\xe9l\xe9s. La page photos est verrouill\xe9e.",            'vote'),
    (u'R\xe9sultats',            u'Classement publi\xe9. Fermez le vote pour y passer.',                                 'resultats'),
])
note(u"Le passage en Vote est bloqu\xe9 tant qu’il reste des photos sans membre ou sans th\xe8me assign\xe9.")
h3(u'G\xe9rer les photos')
bullet(u"Acc\xe9dez \xe0 la page G\xe9rer les photos → depuis la section Gestion.")
bullet(u"Uploadez plusieurs fichiers en une fois. Pr\xe9-assignez membre et th\xe8me avant d’uploader.")
bullet(u"Modifiez l’assignation membre/th\xe8me directement sous chaque vignette.")
bullet(u"La barre de progression indique combien de membres ont une photo par th\xe8me.")

h2(u'3.8 Exposition')
h3(u'Cr\xe9er')
numbered(u"S\xe9lectionnez le type Exposition.")
numbered(u"Renseignez le titre et la description.")
numbered(u"D\xe9finissez les dates d’id\xe9ation (d\xe9but et fin, obligatoires).")
numbered(u"Optionnel : date physique de l’exposition et date d’ouverture au public.")
numbered(u"D\xe9finissez le nombre max de photos par membre.")
numbered(u"Cliquez Cr\xe9er — l’exposition d\xe9marre en phase Id\xe9ation.")
h3(u'Cycle de vie')
statuts_table([
    (u'Id\xe9ation',  u"Les membres proposent des th\xe8mes. Bouton Terminer l’id\xe9ation pour passer \xe0 la suite.",  'ideation'),
    (u'Nettoyage',    u"Filtrez et \xe9ditez les th\xe8mes. Configurez le vote et cliquez Mettre en votation.",               'nettoyage'),
    (u'Votation',     u"Les membres votent. Bouton Terminer le vote. Retour au nettoyage possible si ex-\xe6quo.",            'votation'),
    (u'Soumission',   u"Choisissez le th\xe8me d\xe9finitif, les membres uploadent leurs photos. Bouton Cl\xf4turer.",       'soumission'),
    (u'Cl\xf4tur\xe9',u"Galerie des photos visible. Acc\xe8s public selon la date d’ouverture au public.",              'cloture'),
])
note(u"Le titre, la description et les dates sont modifiables \xe0 tout moment via ✏ Modifier.")

page_break()


# ===========================================================================
# 4. POUR LES ADMINS
# ===========================================================================
h1(u'4. Pour les admins')
para(
    u"En plus de tout ce que peut faire un contributeur, l’admin acc\xe8de au panneau "
    u"d’administration via Admin dans la barre de navigation. Il g\xe8re les membres, "
    u"configure le site, \xe9dite les pages statiques et peut intervenir sur toutes les "
    u"activit\xe9s et publications sans restriction."
)

h2(u'4.1 Gestion des membres')
para(u"Accessible via Admin → Membres. Pour chaque membre, vous pouvez :")
bullet(u"Changer le r\xf4le : Membre, Contributeur, Admin.")
bullet(u"Suspendre un compte (le membre ne peut plus se connecter).")
bullet(u"Voir les statistiques de stockage utilis\xe9.")
bullet(u"Inviter un nouveau membre par email (lien d’invitation).")

h2(u'4.2 Tableau de bord')
para(u"Accessible via Admin → Tableau de bord. Vue synth\xe9tique de l’activit\xe9 du site :")
bullet(u"Nombre de membres, photos, activit\xe9s en cours.")
bullet(u"R\xe9partition du stockage Firebase par cat\xe9gorie (portfolio, th\xe8mes, oneshots…).")
bullet(u"Distribution des membres par r\xf4le.")
note(u"Ces chiffres sont calcul\xe9s en temps r\xe9el — ils refl\xe8tent l’\xe9tat actuel de la base de donn\xe9es.")

h2(u'4.3 Maintenance')
para(u"Accessible via Admin → Maintenance. Outils techniques pour corriger des incoh\xe9rences dans la base de donn\xe9es :")
bullet(u"Recalculer le stockage — resynchronise les compteurs de stockage de tous les membres en cas de valeurs erron\xe9es.")
bullet(u"Recalculer les compteurs — recalcule le nombre de photos par membre en cas de valeur incorrecte.")
note(
    u"IMPORTANT : Ces op\xe9rations lisent toute la base de donn\xe9es et ne doivent \xeatre "
    u"utilis\xe9es qu’en cas de probl\xe8me av\xe9r\xe9. Elles ne modifient pas les photos ni "
    u"les donn\xe9es des membres — elles corrigent uniquement les compteurs d\xe9riv\xe9s. "
    u"A n’utiliser qu’exceptionnellement."
)

h2(u'4.4 Configuration du site')
para(u"Accessible via Admin → Configuration.")
fields_table([
    (u'Image de la page d’accueil',
     u"Photo affich\xe9e en fond sur la page d’accueil. Peut \xeatre une image fixe choisie "
     u"manuellement, ou automatiquement la photo gagnante du dernier th\xe8me du mois."),
    (u'Cat\xe9gories de photos',
     u"Les cat\xe9gories disponibles lors de l’upload d’une photo dans un portfolio."),
])

h2(u'4.5 Pages du site')
para(u"Accessible via Admin → Pages du site. \xc9ditez le contenu des pages statiques :")
bullet(u"Histoire — page « Notre histoire » dans la section Le Club.")
bullet(u"Bureau — pr\xe9sentation des membres du bureau.")
bullet(u"Adh\xe9sion — informations et tarifs d’adh\xe9sion.")
bullet(u"Contact — texte de la page de contact.")
bullet(u"Charte — charte d’utilisation du site. Incr\xe9menter la version oblige tous les membres \xe0 la relire et accepter.")
bullet(u"CGU — conditions g\xe9n\xe9rales d’utilisation. M\xeame logique que la charte.")
bullet(u"Mentions l\xe9gales / Confidentialit\xe9 — pages l\xe9gales obligatoires.")
note(u"Le contenu est en Markdown. Un aper\xe7u en temps r\xe9el est disponible dans l’\xe9diteur.")

h2(u'4.6 Dossiers documents')
para(
    u"Accessible via Admin → Dossiers documents. Cr\xe9ez, renommez et r\xe9ordonnez les "
    u"dossiers de l’espace Documents. La suppression d’un dossier ne supprime pas "
    u"les fichiers qu’il contenait."
)


# ---------------------------------------------------------------------------
out = 'guide_du_site.docx'
doc.save(out)
print(u'OK  ' + out + ' genere.')
